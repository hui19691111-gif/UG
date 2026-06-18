from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


BASE = Path(r"G:\UG6AF7~1\features\12_BIA~1\DOC_BU~1")
OUT = BASE / "roller_exposure_customer_demo.docx"


def set_run_font(run, size=None, bold=None, color=None, font="Microsoft YaHei"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(paragraph, size=11, color="222222", font="Microsoft YaHei"):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color, font=font)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color="222222", size=10.5, fill=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if fill:
        shade_cell(cell, fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)
                set_cell_margins(row.cells[idx])


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(title)
    set_run_font(r, size=24, bold=True, color="0B2545")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    r = p.add_run(subtitle)
    set_run_font(r, size=12.5, color="4B5563")


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    if level == 1:
        set_run_font(r, size=16, bold=True, color="2E74B5")
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        set_run_font(r, size=13, bold=True, color="2E74B5")
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        set_run_font(r, size=12, bold=True, color="1F4D78")
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, size=11, color="222222")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="222222")
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="222222")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=9, color="555555")


def add_image(doc, path, caption, width=6.35):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    add_caption(doc, caption)


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.3])
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F6F9")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color="0B2545")
    p.add_run("\n")
    r = p.add_run(text)
    set_run_font(r, size=10.5, color="222222")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], h, bold=True, color="0B2545", fill="F4F6F9", align=WD_ALIGN_PARAGRAPH.CENTER)
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            align = WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(row.cells[idx], str(value), size=10, align=align)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.2

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("辊筒精细图形曝光设备客户演示方案")
    set_run_font(r, size=9, color="666666")


def build():
    doc = Document()
    setup_styles(doc)

    add_title(
        doc,
        "辊筒精细图形曝光设备客户演示方案",
        "一期单光机头验证，二期同平台扩展至四光机头 | 加工长度 500 mm -> 1000 mm",
    )
    add_body(
        doc,
        "我司基于当前结构示意和技术要求，整理本方案用于客户前期技术沟通和方案演示。我们建议先配置 1 个 2 μm 级光机头完成 500 mm 工件验证，后续在同一机台架构上扩展为 4 个光机头、光机间距 250 mm、适配 1000 mm 工件长度。",
    )
    add_callout(
        doc,
        "我司建议的沟通口径",
        "我们不建议首台机一次性满配，而是按“平台能力先到位、光机数量分阶段配置”的方式推进：一期先验证核心曝光、对焦、旋转同步和工艺窗口；二期再在预留平台上增加光机头，从而降低重复设计风险，并保护客户前期投入。",
    )
    add_table(
        doc,
        ["阶段", "配置目标", "工件能力", "客户价值"],
        [
            ["一期", "我司先配置 1 个 2 μm 级光机头；完成单头曝光验证；预留四头安装平台", "加工长度 500 mm；直径建议覆盖 Φ90-Φ220 mm", "帮助客户先验证核心工艺、降低首台投入、缩短样机交付周期"],
            ["二期", "我司在同一平台增加至 4 个光机头；光机中心距 250 mm；建立多头拼接校正", "加工长度扩展至 1000 mm；同一装夹旋转平台延用", "帮助客户提升加工效率与幅面能力，减少二次开发和整机重做"],
        ],
        [0.9, 2.15, 1.55, 1.7],
    )
    add_image(doc, BASE / "machine_overview.png", "图 1  四光机头扩展状态三维示意，用于说明平台预留和后续扩展能力", width=6.35)

    add_heading(doc, "1. 我司对客户需求的理解", 1)
    add_body(doc, "我司理解客户当前的核心要求可以拆分为两个阶段：首先需要一台可用于工艺验证和样件加工的单光机头设备；后续需要在该机台基础上扩展为四光机头阵列，以覆盖更长工件和更高节拍需求。")
    add_bullet(doc, "一期目标：我司配置 1 个光机头，优先完成 500 mm 长度工件的外圆精细图形曝光。")
    add_bullet(doc, "二期目标：我司在原机台平台上增加至 4 个光机头，光机头间距 250 mm，加工工件长度扩展至 1000 mm。")
    add_bullet(doc, "总体约束：我们将围绕辊筒外圆曝光所需的高精度旋转同步、焦面稳定、光机位姿可调和后续多头拼接补偿能力进行设计。")
    add_bullet(doc, "设计原则：我们在一期即将平台、基座、运动轴、装夹旋转机构和控制系统按二期能力预留，光机头数量则分阶段投入。")

    add_heading(doc, "2. 总体设备方案", 1)
    add_body(doc, "我司方案采用大理石或等效高稳定性基座，工件由固定抓盘和移动抓盘两端支撑，固定端内置高精度 DD 马达作为分度/旋转驱动，移动端沿 X 向调整以适配不同长度工件。我们将光机头安装在高稳定安装平台上，平台预留四头阵列安装位置，并配置对位 CCD、六轴微调和短行程调焦机构。")
    add_image(doc, BASE / "pdf_pages" / "pdf_page_1_cropped.png", "图 2  整机结构示意：旋转抓盘、移动抓盘、CCD 对位、X 向行程和大理石底座", width=6.35)

    add_heading(doc, "3. 分阶段实施方案", 1)
    add_heading(doc, "3.1 一期：单光机头 500 mm 工件验证", 2)
    add_body(doc, "一期我司以“完整平台 + 单头配置”为目标。我们会将机械平台、装夹旋转、光机安装平台、对位和控制系统按后续四头方案进行结构预留；实际交付时先安装 1 个 2 μm 级光机头，用于完成 500 mm 工件的曝光能力验证。")
    add_bullet(doc, "光机配置：我司配置 1 个 2 μm 级光机头，带六轴精密调节与短行程实时调焦能力。")
    add_bullet(doc, "工件范围：我司一期重点覆盖长度 500 mm 工件，直径范围建议按 Φ90-Φ220 mm 设计。")
    add_bullet(doc, "运动与对位：我们保留 X 轴 550 mm 行程、Y 轴 150 mm 调整能力和 CCD 对位能力，为后续多头布局和工件长度扩展预留。")
    add_bullet(doc, "验证重点：我们重点验证曝光光斑、对焦稳定性、旋转同步、图形变形/错位、工件装夹重复性和基础软件流程。")

    add_heading(doc, "3.2 二期：四光机头 1000 mm 工件扩展", 2)
    add_body(doc, "二期我司在一期机台基础上增加光机头数量，形成 4 个光机头阵列，光机中心距按 250 mm 设计。由于平台、安装面和运动行程在一期已预留，二期主要工作集中在光机安装、单头标定、多头拼接补偿和控制软件扩展。")
    add_bullet(doc, "光机阵列：我们配置 4 个光机头，间距 250 mm，覆盖更长轴向曝光区域。")
    add_bullet(doc, "工件长度：我们将加工长度能力由 500 mm 扩展至 1000 mm。")
    add_bullet(doc, "拼接能力：我们建立光机间位置、倍率、旋转角度和能量一致性的标定方法。")
    add_bullet(doc, "控制升级：我们增加多光机同步曝光、分区图形分配、拼接补偿、运行状态监控和参数追溯。")
    add_image(doc, BASE / "pdf_pages" / "pdf_page_2_cropped.png", "图 3  光机安装平台示意：一期安装 1 个光机，平台预留 4 个光机位，间距 250 mm", width=6.35)

    add_heading(doc, "4. 我司建议的核心技术指标", 1)
    add_table(
        doc,
        ["项目", "一期建议指标", "二期扩展指标"],
        [
            ["光机头数量", "1 个", "4 个，光机中心距 250 mm"],
            ["加工工件长度", "500 mm", "1000 mm"],
            ["工件直径范围", "Φ90-Φ220 mm", "沿用一期装夹范围"],
            ["最小光斑/像素单元", "约 1.98 μm x 1.98 μm", "各光机头保持一致性"],
            ["旋转能力", "360° 连续旋转，DD 马达分度/同步", "沿用并配合多头同步曝光"],
            ["光机平台行程", "X 轴 550 mm，Y 轴 150 mm 预留", "满足 4 头安装及长工件覆盖"],
            ["对焦能力", "手动调焦 + 自动对焦，控制精度目标 0.005 mm", "多头独立对焦与焦面一致性校正"],
            ["图形质量目标", "显影后变形、错位目标不大于 0.2 μm", "增加多头拼接误差控制"],
        ],
        [1.65, 2.3, 2.25],
    )

    add_heading(doc, "5. 关键机构说明", 1)
    add_heading(doc, "5.1 工件装夹与旋转机构", 2)
    add_body(doc, "我司在固定抓盘内置高精度 DD 马达，用于提供工件旋转动力和分度控制；移动抓盘用于适配不同长度工件，可沿 X 轴方向移动。我们建议保留 550 mm 行程，并采用高稳定导向或气浮支撑方案。")
    add_heading(doc, "5.2 光机安装平台", 2)
    add_body(doc, "我司建议光机安装平台采用大理石或等效高稳定材料，平台在一期即按 4 个光机安装位进行预留。单光机阶段只安装 1 个光机头，但我们会将平台安装基准、线缆走线、气路/冷却接口和安装孔位按四头状态统一规划。")
    add_heading(doc, "5.3 六轴精密调节与调焦", 2)
    add_body(doc, "我司为每个光机头配置六轴精密调节能力，包括 X、Y、Z、Rx、Ry、Rz 六个自由度。Z 向需同时支持手动调焦和自动对焦，自动对焦用于补偿辊筒直径差异、装夹偏差和旋转过程中的焦面变化。")
    add_heading(doc, "5.4 CCD 对位与标定", 2)
    add_body(doc, "我司配置高清 CCD 组件，用于工件对位拍照、装夹位置确认、图形基准识别和后续多头拼接标定。二期扩展后，我们可利用 CCD 数据建立各光机头之间的坐标转换和误差补偿关系。")

    add_heading(doc, "6. 我司建议的加工流程", 1)
    for item in [
        "装夹工件：固定端夹紧，移动抓盘按工件长度调整位置并锁紧。",
        "工件对位：CCD 拍照识别基准，确认辊筒轴向和周向零位。",
        "调焦与标定：完成光机头焦距、角度、能量和图形坐标标定。",
        "导入图形：软件导入曝光图形，生成轴向/周向曝光路径。",
        "同步曝光：旋转轴、X/Y 运动、光机曝光和对焦机构联动执行。",
        "显影检测：对曝光显影后的图形变形、错位、线宽和拼接区域进行检测。",
    ]:
        add_number(doc, item)

    doc.add_page_break()
    add_heading(doc, "7. 一期验收重点", 1)
    add_table(
        doc,
        ["类别", "验收内容", "说明"],
        [
            ["机械", "500 mm 工件装夹、旋转稳定性、X/Y 运动行程与重复性", "验证平台基础能力是否满足后续扩展"],
            ["光学", "单光机光斑尺寸、能量稳定性、焦面稳定性", "确认 2 μm 级曝光能力"],
            ["对焦", "手动调焦、自动对焦、焦面补偿响应", "目标控制精度 0.005 mm"],
            ["软件", "图形导入、路径生成、参数保存、运行记录", "为二期多头同步做软件架构预留"],
            ["工艺", "样件显影后的图形变形、错位和重复性", "作为二期扩展前的核心放行条件"],
        ],
        [1.0, 2.8, 2.1],
    )

    add_heading(doc, "8. 我司建议二期扩展前确认事项", 1)
    add_bullet(doc, "四光机头的最终曝光幅面分配方式，以及每个光机头负责的轴向加工区间。")
    add_bullet(doc, "光机头中心距 250 mm 与实际有效曝光宽度、拼接重叠量之间的关系。")
    add_bullet(doc, "多头之间的光斑一致性、能量一致性、焦面一致性和坐标转换标定方法。")
    add_bullet(doc, "1000 mm 工件在旋转过程中的跳动、挠曲、热稳定性和夹持同轴度控制要求。")
    add_bullet(doc, "最终验收图形、检测方法、检测设备分辨率和判定标准。")

    doc.add_page_break()
    add_heading(doc, "9. 方案价值总结", 1)
    add_body(doc, "我司方案的优势在于把首台设备风险控制在单光机头验证阶段，同时让机台基础能力一次做到位。客户可先用 500 mm 工件完成工艺验证、样件验证和软件流程验证；当工艺窗口确认后，我们可在原平台上增加至 4 个光机头，实现 1000 mm 工件加工和更高生产效率。")
    add_callout(
        doc,
        "我司对客户的核心承诺",
        "我们的一期交付可验证，二期扩展不推倒重来。机台平台、装夹旋转、光机安装面、运动行程和控制系统均围绕四头扩展预留，从而降低后续升级风险。",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
