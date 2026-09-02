from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "manual_assets"
OUT = DOCS / "批量转钣金使用说明.docx"


def font_path(name: str = "msyh.ttc") -> str:
    return str(Path("C:/Windows/Fonts") / name)


FONT = font_path("msyh.ttc")
FONT_BOLD = font_path("msyhbd.ttc")


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def rounded_rect(draw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered(draw, box, text, font, fill):
    x1, y1, x2, y2 = box
    w, h = text_size(draw, text, font)
    draw.text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2 - 2), text, font=font, fill=fill)


def create_flow_image(path: Path):
    img = Image.new("RGB", (1400, 760), "white")
    d = ImageDraw.Draw(img)
    title = ImageFont.truetype(FONT_BOLD, 42)
    h = ImageFont.truetype(FONT_BOLD, 30)
    body = ImageFont.truetype(FONT, 24)
    small = ImageFont.truetype(FONT, 21)

    d.text((50, 35), "批量转钣金推荐操作流程", font=title, fill=(20, 47, 82))
    steps = [
        ("1", "打开功能", "装配或单部件环境\n启动“自动转钣金”"),
        ("2", "筛选对象", "在列表中用关键字、\n图层和属性条件缩小范围"),
        ("3", "设置参数", "确认固定面方向、\n颜色、避让槽和过滤项"),
        ("4", "手动指定", "需要时点击“手动选择...”\n选择基面和X向边"),
        ("5", "执行处理", "点击确定，程序处理\n未手动处理的对象"),
        ("6", "查看结果", "全部成功不弹框；\n失败时弹出明细"),
    ]
    w = 380
    hbox = 190
    x0 = 70
    y0 = 145
    col_gap = 65
    row_gap = 70
    for pos, (idx, head, desc) in enumerate(steps):
        col = pos % 3
        row = pos // 3
        x = x0 + col * (w + col_gap)
        y = y0 + row * (hbox + row_gap)
        box = (x, y, x + w, y + hbox)
        rounded_rect(d, box, 22, (238, 244, 251), (119, 154, 190), 3)
        rounded_rect(d, (x + 18, y + 18, x + 64, y + 64), 23, (46, 117, 182))
        draw_centered(d, (x + 18, y + 18, x + 64, y + 64), idx, h, "white")
        d.text((x + 84, y + 24), head, font=h, fill=(20, 47, 82))
        d.multiline_text((x + 28, y + 86), desc, font=body, fill=(45, 55, 72), spacing=8)
        if pos not in (2, 5):
            ax = x + w + 10
            ay = y + 95
            d.line((ax, ay, ax + col_gap - 24, ay), fill=(72, 96, 122), width=4)
            d.polygon([(ax + col_gap - 24, ay - 9), (ax + col_gap - 10, ay), (ax + col_gap - 24, ay + 9)], fill=(72, 96, 122))
        if pos == 2:
            ax = x + w / 2
            d.line((ax, y + hbox + 12, ax, y + hbox + row_gap - 18), fill=(72, 96, 122), width=4)
            d.polygon([(ax - 9, y + hbox + row_gap - 18), (ax, y + hbox + row_gap - 4), (ax + 9, y + hbox + row_gap - 18)], fill=(72, 96, 122))
    d.text((75, 670), "提示：手动选择只处理你指定的体；返回主对话框后点击“确定”，其余未处理对象继续按自动方式执行。", font=small, fill=(96, 68, 0))
    img.save(path)


def create_manual_image(path: Path):
    img = Image.new("RGB", (1400, 720), "white")
    d = ImageDraw.Draw(img)
    title = ImageFont.truetype(FONT_BOLD, 42)
    h = ImageFont.truetype(FONT_BOLD, 29)
    body = ImageFont.truetype(FONT, 24)
    small = ImageFont.truetype(FONT, 21)

    d.text((50, 35), "手动选择基面与X向边", font=title, fill=(20, 47, 82))
    d.text((55, 95), "适合自动方向不符合要求、标记线方向需要人工确认、或个别零件需要特殊展开方向的场景。", font=body, fill=(60, 70, 86))

    panel = (80, 165, 780, 620)
    rounded_rect(d, panel, 24, (245, 248, 252), (160, 174, 190), 2)
    # part sketch
    top = [(245, 250), (590, 210), (700, 315), (350, 360)]
    side = [(350, 360), (700, 315), (700, 450), (350, 505)]
    front = [(245, 250), (350, 360), (350, 505), (245, 385)]
    d.polygon(side, fill=(216, 226, 235), outline=(97, 118, 138))
    d.polygon(front, fill=(198, 216, 229), outline=(97, 118, 138))
    d.polygon(top, fill=(140, 194, 126), outline=(69, 116, 61))
    d.line((365, 358, 665, 320), fill=(207, 61, 54), width=9)
    d.polygon([(665, 320), (640, 305), (647, 337)], fill=(207, 61, 54))
    d.text((232, 215), "基面", font=h, fill=(37, 105, 48))
    d.text((520, 285), "X向边", font=h, fill=(176, 36, 31))

    notes = [
        ("先选基面", "选择需要作为展开固定面的平面。"),
        ("再选X向边", "边必须在刚选的基面上，并且是直线边。"),
        ("可重复选择", "一个体完成后可继续选下一个体；取消则返回主对话框。"),
        ("未手动的对象", "回到主对话框点击确定后，按自动方式继续处理。"),
    ]
    x, y = 840, 170
    for i, (head, desc) in enumerate(notes, 1):
        rounded_rect(d, (x, y, x + 480, y + 95), 18, (238, 244, 251), (184, 199, 215), 2)
        rounded_rect(d, (x + 18, y + 25, x + 60, y + 67), 21, (46, 117, 182))
        draw_centered(d, (x + 18, y + 25, x + 60, y + 67), str(i), h, "white")
        d.text((x + 78, y + 16), head, font=h, fill=(20, 47, 82))
        d.text((x + 78, y + 55), desc, font=small, fill=(55, 66, 82))
        y += 112

    d.text((90, 640), "注意：手动选择面时，实体需要具备名称为 sulian 且有值的属性。", font=small, fill=(120, 55, 20))
    img.save(path)


def create_result_image(path: Path):
    img = Image.new("RGB", (1400, 520), "white")
    d = ImageDraw.Draw(img)
    title = ImageFont.truetype(FONT_BOLD, 42)
    h = ImageFont.truetype(FONT_BOLD, 28)
    body = ImageFont.truetype(FONT, 23)
    small = ImageFont.truetype(FONT, 20)

    d.text((50, 35), "结果提示方式", font=title, fill=(20, 47, 82))
    cards = [
        ((70, 135, 655, 420), (232, 248, 236), (61, 135, 80), "全部成功", "不弹最终结果框", "处理完后保持界面安静，避免每次成功都需要手动关闭弹窗。"),
        ((745, 135, 1330, 420), (254, 239, 239), (176, 52, 45), "有失败或异常", "弹出结果明细", "可看到转钣金成功数、展开成功数、失败数量、图层和失败原因。"),
    ]
    for box, fill, accent, head, sub, desc in cards:
        rounded_rect(d, box, 22, fill, accent, 3)
        d.text((box[0] + 34, box[1] + 34), head, font=h, fill=accent)
        d.text((box[0] + 34, box[1] + 82), sub, font=h, fill=(34, 46, 62))
        d.multiline_text((box[0] + 34, box[1] + 145), desc, font=body, fill=(55, 66, 82), spacing=8)
    d.text((70, 455), "过滤掉的对象不等于失败；例如不满足材料、数量、图层等条件的对象会被跳过。", font=small, fill=(92, 70, 35))
    img.save(path)


def badge(draw, center, text, font):
    x, y = center
    draw.ellipse((x - 22, y - 22, x + 22, y + 22), fill=(221, 71, 64), outline="white", width=3)
    draw_centered(draw, (x - 22, y - 22, x + 22, y + 22), text, font, "white")


def create_main_dialog_image(path: Path):
    img = Image.new("RGB", (1400, 950), "white")
    d = ImageDraw.Draw(img)
    title = ImageFont.truetype(FONT_BOLD, 42)
    h = ImageFont.truetype(FONT_BOLD, 25)
    body = ImageFont.truetype(FONT, 22)
    small = ImageFont.truetype(FONT, 18)
    num = ImageFont.truetype(FONT_BOLD, 24)

    d.text((50, 35), "主对话框控件编号", font=title, fill=(20, 47, 82))
    dlg = (100, 105, 1300, 900)
    rounded_rect(d, dlg, 20, (246, 248, 251), (122, 142, 164), 3)
    d.rectangle((100, 105, 1300, 155), fill=(35, 77, 120))
    d.text((125, 116), "自动转钣金", font=h, fill="white")

    # Strategy group
    rounded_rect(d, (135, 185, 1265, 440), 14, (255, 255, 255), (190, 202, 216), 2)
    d.text((155, 198), "系数策略", font=h, fill=(31, 77, 120))
    controls = [
        (1, (170, 250, 430, 285), "按标记线找基面", "toggle"),
        (2, (170, 300, 430, 335), "运行完成自动保存", "toggle"),
        (3, (170, 350, 430, 385), "手动选择...", "button"),
        (4, (500, 245, 815, 290), "固定面方向：下折多 ▼", "combo"),
        (5, (500, 310, 815, 355), "失败处理：隐藏成功项 ▼", "combo"),
        (6, (500, 375, 815, 420), "展开基面颜色", "toggle"),
        (7, (840, 375, 960, 420), "颜色", "color"),
        (8, (1005, 245, 1225, 290), "规则设置...", "button"),
    ]
    for n, box, text, kind in controls:
        fill = (239, 246, 253) if kind in ("combo", "button") else (255, 255, 255)
        rounded_rect(d, box, 8, fill, (132, 160, 190), 2)
        if kind == "toggle":
            d.rectangle((box[0] + 12, box[1] + 9, box[0] + 34, box[1] + 31), outline=(90, 116, 145), width=2)
            d.text((box[0] + 45, box[1] + 6), text, font=body, fill=(35, 47, 62))
        elif kind == "color":
            d.rectangle((box[0] + 18, box[1] + 10, box[0] + 58, box[1] + 35), fill=(238, 199, 70), outline=(90, 90, 90))
            d.text((box[0] + 70, box[1] + 7), text, font=body, fill=(35, 47, 62))
        else:
            draw_centered(d, box, text, body, (35, 47, 62))
        badge(d, (box[0] - 8, box[1] + 2), str(n), num)

    # Parameters
    rounded_rect(d, (135, 465, 1265, 595), 14, (255, 255, 255), (190, 202, 216), 2)
    d.text((155, 478), "基础参数", font=h, fill=(31, 77, 120))
    params = [
        (9, (180, 530, 335, 570), "避让槽深"),
        (10, (455, 530, 610, 570), "避让槽宽"),
        (11, (730, 530, 885, 570), "内R半径"),
    ]
    for n, box, label in params:
        d.text((box[0] - 95, box[1] + 8), label, font=body, fill=(35, 47, 62))
        rounded_rect(d, box, 6, (255, 255, 255), (132, 160, 190), 2)
        d.text((box[0] + 12, box[1] + 8), "0.2", font=body, fill=(80, 90, 105))
        badge(d, (box[0] - 10, box[1] - 4), str(n), num)

    # Advanced
    rounded_rect(d, (135, 620, 1265, 835), 14, (255, 255, 255), (190, 202, 216), 2)
    d.text((155, 633), "高级过滤", font=h, fill=(31, 77, 120))
    adv = [
        (12, (170, 690, 430, 725), "跳过过小实体", "toggle"),
        (13, (170, 750, 335, 790), "最小长度", "input"),
        (14, (450, 750, 615, 790), "最小宽度", "input"),
        (15, (500, 690, 835, 725), "跳过螺母、螺柱、螺钉", "toggle"),
        (16, (875, 690, 1205, 725), "按图层范围找实体", "toggle"),
        (17, (875, 750, 1030, 790), "起始图层", "input"),
        (18, (1090, 750, 1245, 790), "结束图层", "input"),
        (19, (170, 805, 500, 835), "每层只处理最大实体", "toggle"),
    ]
    for n, box, text, kind in adv:
        if kind == "toggle":
            d.rectangle((box[0] + 12, box[1] + 7, box[0] + 34, box[1] + 29), outline=(90, 116, 145), width=2)
            d.text((box[0] + 45, box[1] + 4), text, font=body, fill=(35, 47, 62))
        else:
            d.text((box[0] - 90, box[1] + 8), text, font=body, fill=(35, 47, 62))
            rounded_rect(d, box, 6, (255, 255, 255), (132, 160, 190), 2)
        badge(d, (box[0] - 8, box[1] + 2), str(n), num)

    d.text((155, 860), "说明：编号只用于本文档讲解，实际软件界面不显示这些红色编号。", font=small, fill=(96, 68, 0))
    img.save(path)


def create_picker_dialog_image(path: Path):
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    title = ImageFont.truetype(FONT_BOLD, 42)
    h = ImageFont.truetype(FONT_BOLD, 24)
    body = ImageFont.truetype(FONT, 21)
    small = ImageFont.truetype(FONT, 17)
    num = ImageFont.truetype(FONT_BOLD, 22)

    d.text((50, 35), "装配列表窗口控件编号", font=title, fill=(20, 47, 82))
    dlg = (70, 105, 1430, 835)
    rounded_rect(d, dlg, 18, (246, 248, 251), (122, 142, 164), 3)
    d.rectangle((70, 105, 1430, 155), fill=(35, 77, 120))
    d.text((95, 116), "Batch sheet metal parts / 批量选择部件", font=h, fill="white")

    # Top filters
    top_controls = [
        (1, (135, 185, 365, 225), "过滤关键字"),
        (2, (410, 185, 560, 225), "部件图层 ▼"),
        (3, (585, 185, 735, 225), "图层范围"),
        (4, (760, 185, 850, 225), "清空"),
    ]
    for n, box, text in top_controls:
        rounded_rect(d, box, 7, (255, 255, 255), (132, 160, 190), 2)
        draw_centered(d, box, text, body, (35, 47, 62))
        badge(d, (box[0] - 8, box[1] + 2), str(n), num)

    checks = [
        (5, (135, 245), "有材料属性值"),
        (6, (335, 245), "有数量属性值"),
        (7, (535, 245), "部件不含钣金特征"),
        (8, (785, 245), "显示的部件"),
    ]
    for n, (x, y), text in checks:
        d.rectangle((x, y, x + 24, y + 24), outline=(90, 116, 145), width=2)
        d.text((x + 34, y - 2), text, font=body, fill=(35, 47, 62))
        badge(d, (x - 10, y - 4), str(n), num)

    # List
    list_box = (110, 305, 1390, 710)
    d.rectangle(list_box, fill="white", outline=(132, 160, 190), width=2)
    columns = ["选择", "部件名称", "材料", "数量", "图层", "自定义列1", "自定义列2"]
    widths = [90, 330, 170, 120, 120, 220, 220]
    x = list_box[0]
    for col, width in zip(columns, widths):
        d.rectangle((x, list_box[1], x + width, list_box[1] + 42), fill=(232, 238, 245), outline=(132, 160, 190))
        draw_centered(d, (x, list_box[1], x + width, list_box[1] + 42), col, small, (20, 47, 82))
        x += width
    for row in range(1, 7):
        y = list_box[1] + 42 * row
        d.line((list_box[0], y, list_box[2], y), fill=(220, 226, 234), width=1)
        d.text((130, y + 9), "☑", font=body, fill=(35, 47, 62))
        d.text((220, y + 9), f"part_{row:02d}.prt", font=small, fill=(55, 66, 82))
        d.text((545, y + 9), "SPCC", font=small, fill=(55, 66, 82))
        d.text((720, y + 9), str(row), font=small, fill=(55, 66, 82))
        d.text((840, y + 9), str(10 + row), font=small, fill=(55, 66, 82))
    badge(d, (100, 300), "9", num)
    d.text((125, 720), "9  部件列表：勾选需要处理的部件；属性列可通过上方下拉选择显示内容。", font=small, fill=(96, 68, 0))

    buttons = [
        (10, (125, 770, 225, 815), "全选"),
        (11, (245, 770, 365, 815), "全不选"),
        (12, (1110, 770, 1210, 815), "确定"),
        (13, (1230, 770, 1330, 815), "取消"),
    ]
    for n, box, text in buttons:
        rounded_rect(d, box, 8, (239, 246, 253), (132, 160, 190), 2)
        draw_centered(d, box, text, body, (35, 47, 62))
        badge(d, (box[0] - 8, box[1] + 2), str(n), num)

    img.save(path)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def style_table(table, header_fill="E8EEF5"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True


def compact_table_text(table, size=9):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.size = Pt(size)
                    r.font.name = "微软雅黑"
                    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_cropped_real_dialog_capture(doc: Document, path: Path):
    """Insert the user-provided NX capture and crop it in Word without redrawing the UI."""
    shape = doc.add_picture(str(path))
    shape.width = Inches(3.05)
    shape.height = Inches(5.72)

    blip_fill = shape._inline.graphic.graphicData.pic.blipFill
    src_rect = OxmlElement("a:srcRect")
    # Original image is 750 x 844. Keep the real dialog at approximately
    # x=112..374, y=173..653 and hide the unrelated chat/error area.
    src_rect.set("l", "14933")
    src_rect.set("t", "20498")
    src_rect.set("r", "50133")
    src_rect.set("b", "22630")
    blip_fill.insert(1, src_rect)

    paragraph = shape._inline.getparent()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_margins(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)


def set_style(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2
    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        s = styles[style_name]
        s.font.name = "微软雅黑"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True
        s.paragraph_format.space_before = Pt(10 if style_name != "Heading 1" else 14)
        s.paragraph_format.space_after = Pt(5)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("批量转钣金功能使用说明")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(20, 47, 82)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("适用于 UG 智辉钣金插件 - 自动转钣金/批量展开")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(92, 105, 122)


def add_note(doc: Document, title: str, text: str, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10.5)
    r.font.bold = True
    r.font.color.rgb = RGBColor(31, 77, 120)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(10)


def add_bullets(doc: Document, items: list[str]):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(text)
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_numbered(doc: Document, items: list[str]):
    for index, text in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{index}.  {text}")
        r.font.name = "微软雅黑"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def add_option_table(doc: Document):
    rows = [
        ("按标记线找基面", "有标记线时使用。勾选后，展开基面会优先按标记线所在侧来确定方向。"),
        ("运行完成自动保存", "处理完成后自动保存部件。批量处理正式数据时建议开启。"),
        ("手动选择...", "对指定实体人工选择基面和 X 向边。适合个别零件自动方向不符合要求时使用。"),
        ("固定面方向", "选择上折多或下折多，用于控制自动展开时的固定面方向偏好。"),
        ("转钣金展开失败处理方式", "设置失败对象的处理方式，例如隐藏成功项、失败改红或高亮显示。"),
        ("展开基面颜色", "开启后，展开基面按所选颜色着色，便于检查基面是否正确。"),
        ("避让槽深/宽、内R半径", "转钣金时使用的基础工艺参数。"),
        ("跳过过小实体", "按最小长度和最小宽度过滤小实体。"),
        ("跳过螺母、螺柱、螺钉", "批量处理时排除紧固件类对象。"),
        ("按图层范围找实体", "只处理起始图层到结束图层范围内的实体。"),
        ("每层只处理最大实体", "同一图层有多个实体时，只处理最大实体。"),
        ("规则设置...", "打开折弯规则表，维护材料、板厚、角度和 K 因子等规则。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), "选项", True)
    set_cell_text(table.cell(0, 1), "用途", True)
    for name, usage in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], name)
        set_cell_text(cells[1], usage)
    style_table(table)


def add_filter_table(doc: Document):
    rows = [
        ("过滤关键字", "按部件名称或列表中显示的属性内容过滤。"),
        ("图层模式", "可选择“部件图层”或“装配图层”。"),
        ("图层范围", "支持单个图层、逗号分隔和区间输入，例如 10、10,12,15、10-15。"),
        ("有材料属性值", "只显示材料属性有值的部件。"),
        ("有数量属性值", "只显示数量属性有值的部件。"),
        ("部件不含钣金特征", "只显示还没有钣金特征的部件。"),
        ("显示的部件", "只显示当前可见部件。"),
        ("全选/全不选", "只作用于当前过滤后列表中显示的对象。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), "过滤项", True)
    set_cell_text(table.cell(0, 1), "说明", True)
    for name, usage in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], name)
        set_cell_text(cells[1], usage)
    style_table(table)


def add_main_dialog_control_table(doc: Document):
    rows = [
        ("1", "按标记线找基面", "开关", "有标记线要求时勾选。用于让展开方向参考标记线所在侧。"),
        ("2", "运行完成自动保存", "开关", "勾选后处理结束自动保存部件；正式批量处理时建议开启。"),
        ("3", "手动选择...", "按钮", "进入手动选择流程，给指定实体选择基面和 X 向边。"),
        ("4", "固定面方向", "下拉框", "选择“上折多”或“下折多”，用于控制自动展开时的方向偏好。"),
        ("5", "转钣金展开失败处理方式", "下拉框", "设置失败时如何处理对象，例如隐藏成功项、失败改红或高亮。"),
        ("6", "展开基面颜色", "开关", "勾选后给展开基面着色，便于检查选到的基面。"),
        ("7", "颜色选择", "颜色框", "选择展开基面显示的颜色。只有“展开基面颜色”开启时生效。"),
        ("8", "规则设置...", "按钮", "打开折弯规则表，维护材料、板厚、角度和 K 因子。"),
        ("9", "避让槽深", "数字输入框", "输入转钣金时使用的避让槽深度。"),
        ("10", "避让槽宽", "数字输入框", "输入转钣金时使用的避让槽宽度。"),
        ("11", "内R半径", "数字输入框", "输入内 R 半径。没有特殊要求时使用默认值。"),
        ("12", "跳过过小实体", "开关", "勾选后，按最小长度和最小宽度过滤小实体。"),
        ("13", "最小长度", "数字输入框", "和“跳过过小实体”配合使用，小于该长度的实体会被跳过。"),
        ("14", "最小宽度", "数字输入框", "和“跳过过小实体”配合使用，小于该宽度的实体会被跳过。"),
        ("15", "跳过螺母、螺柱、螺钉", "开关", "勾选后批量处理时排除紧固件类对象。"),
        ("16", "按图层范围找实体", "开关", "勾选后只处理指定图层范围内的实体。"),
        ("17", "起始图层", "整数输入框", "输入图层范围的起始层。"),
        ("18", "结束图层", "整数输入框", "输入图层范围的结束层。"),
        ("19", "每层只处理最大实体", "开关", "同一图层有多个实体时，只处理该层最大实体。"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_col_widths(table, [Inches(0.45), Inches(1.45), Inches(0.85), Inches(3.75)])
    headers = ["编号", "控件名称", "类型", "使用说明"]
    for i, text in enumerate(headers):
        set_cell_text(table.cell(0, i), text, True, 9)
    for row in rows:
        cells = table.add_row().cells
        set_table_col_widths(table, [Inches(0.45), Inches(1.45), Inches(0.85), Inches(3.75)])
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, False, 9)
    style_table(table)
    compact_table_text(table, 9)


def add_picker_dialog_control_table(doc: Document):
    rows = [
        ("1", "过滤关键字", "输入框", "输入部件名或属性关键字，列表只显示匹配对象。"),
        ("2", "图层模式", "下拉框", "选择按“部件图层”过滤，或按“装配图层”过滤。"),
        ("3", "图层范围", "输入框", "输入要显示的图层。支持 10、10,12,15、10-15 这种格式。"),
        ("4", "清空", "按钮", "清除关键字、图层范围和过滤开关，恢复列表显示。"),
        ("5", "有材料属性值", "开关", "只显示材料属性有值的部件。"),
        ("6", "有数量属性值", "开关", "只显示数量属性有值的部件。"),
        ("7", "部件不含钣金特征", "开关", "只显示还没有钣金特征的部件。"),
        ("8", "显示的部件", "开关", "只显示当前在装配中可见的部件。"),
        ("9", "部件列表", "列表", "勾选要处理的部件；可查看部件名、材料、数量、图层等信息。"),
        ("10", "全选", "按钮", "勾选当前过滤后列表中显示的全部对象。"),
        ("11", "全不选", "按钮", "取消当前过滤后列表中显示对象的勾选。"),
        ("12", "确定", "按钮", "确认本次列表选择，进入主参数对话框或继续执行。"),
        ("13", "取消", "按钮", "退出列表选择，不执行本次批量处理。"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_col_widths(table, [Inches(0.45), Inches(1.45), Inches(0.85), Inches(3.75)])
    headers = ["编号", "控件名称", "类型", "使用说明"]
    for i, text in enumerate(headers):
        set_cell_text(table.cell(0, i), text, True, 9)
    for row in rows:
        cells = table.add_row().cells
        set_table_col_widths(table, [Inches(0.45), Inches(1.45), Inches(0.85), Inches(3.75)])
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, False, 9)
    style_table(table)
    compact_table_text(table, 9)


def add_rules_dialog_control_table(doc: Document):
    rows = [
        ("选择规则", "下拉框", "选择要维护的折弯角度区间：90°、0～90°、90～180°或180～360°。"),
        ("普通折弯", "下拉框", "指定普通折弯读取扣除1～3、K因子1～3或A1～A3中的哪一列。"),
        ("多刀折圆", "下拉框", "指定多刀折圆读取的系数列。截图中选择的是“K因子3”。"),
        ("多刀折圆最小半径", "数字输入框", "大于0时启用；达到该半径门槛的圆弧折弯按多刀折圆规则取值。单位跟随零件单位。"),
        ("材料分页", "下拉框", "选择“全部”或单一材料，控制下方表格显示范围。"),
        ("新增厚度", "按钮", "为材料新增一个厚度规格，再编辑厚度和各系数列。"),
        ("系数表格", "可编辑表格", "双击单元格维护材料、厚度、扣除值、K因子和A1～A3。只有上方选中的列参与对应计算。"),
        ("保存", "按钮", "将当前规则写入配置文件；每次修改后建议立即保存。"),
        ("导入/导出EXCEL数据", "按钮", "批量交换CSV格式的表格数据。导入后先核对内容，再点击保存。"),
        ("计算K因子", "按钮", "打开K因子计算工具，用试折或已知展开数据辅助计算。"),
        ("确定 / 取消", "按钮", "确定会保存后关闭；取消会放弃尚未保存的修改。"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    set_table_col_widths(table, [Inches(1.6), Inches(1.0), Inches(4.0)])
    headers = ["控件", "类型", "使用说明"]
    for i, text in enumerate(headers):
        set_cell_text(table.cell(0, i), text, True, 9)
    for row in rows:
        cells = table.add_row().cells
        set_table_col_widths(table, [Inches(1.6), Inches(1.0), Inches(4.0)])
        for i, text in enumerate(row):
            set_cell_text(cells[i], text, False, 9)
    style_table(table)
    compact_table_text(table, 9)


def build_doc():
    ASSETS.mkdir(parents=True, exist_ok=True)
    create_flow_image(ASSETS / "流程图.png")
    create_manual_image(ASSETS / "手动选择示意.png")
    create_result_image(ASSETS / "结果提示.png")
    real_dialog_capture = ASSETS / "主对话框真实截图.png"
    if not real_dialog_capture.is_file():
        raise FileNotFoundError(f"缺少真实对话框截图：{real_dialog_capture}")
    assembly_capture = ASSETS / "装配选择真实截图.png"
    if not assembly_capture.is_file():
        raise FileNotFoundError(f"缺少真实装配选择截图：{assembly_capture}")
    rules_capture = ASSETS / "规则表真实截图.png"
    if not rules_capture.is_file():
        raise FileNotFoundError(f"缺少真实规则表截图：{rules_capture}")

    doc = Document()
    set_margins(doc)
    set_style(doc)
    add_title(doc)

    doc.add_heading("1. 功能用途", level=1)
    doc.add_paragraph("批量转钣金用于把选中的实体或装配子部件批量转换为钣金，并自动创建展开结果。它适合多个零件连续处理，也支持对个别零件手动指定展开基面和 X 向。")
    add_note(doc, "使用原则", "正常情况下直接按自动流程执行；只有个别零件方向不符合要求时，再使用“手动选择...”单独指定。")
    doc.add_picture(str(ASSETS / "流程图.png"), width=Inches(6.7))

    doc.add_heading("2. 基本使用步骤", level=1)
    add_numbered(doc, [
        "在 UG 中打开需要处理的单个部件或装配。",
        "启动“自动转钣金”功能。",
        "如果在装配环境运行，先在列表中筛选并勾选要处理的部件。",
        "在主对话框中确认固定面方向、避让槽、内 R、图层范围等参数。",
        "需要人工控制方向的实体，点击“手动选择...”并按提示选择基面和 X 向边。",
        "返回主对话框后点击“确定”，程序会继续处理未手动处理的对象。",
        "全部成功时不弹最终提示；有失败或异常时，查看弹框中的失败明细。",
    ])

    doc.add_heading("3. 主对话框界面说明", level=1)
    doc.add_paragraph("下图来自 NX 中实际运行的“自动转钣金”对话框。截图右侧的错误窗口和聊天背景已通过文档裁切隐藏，没有重画或替换任何对话框控件。请按界面从上到下，对照下表查看每个选项的用途。")
    add_cropped_real_dialog_capture(doc, real_dialog_capture)
    add_main_dialog_control_table(doc)

    doc.add_heading("4. 装配列表界面说明", level=1)
    doc.add_paragraph("在装配环境下运行时，会先出现待处理部件列表。下图来自实际装配模型运行界面。先设置过滤条件，再勾选需要处理的部件；点击“确定”后才进入主参数对话框。")
    doc.add_picture(str(assembly_capture), width=Inches(6.7))
    add_picker_dialog_control_table(doc)
    add_note(doc, "图层输入示例", "输入 10 表示只看 10 层；输入 10,12,15 表示只看这几个图层；输入 10-15 表示 10 到 15 层。")

    doc.add_heading("5. 折弯系数规则表", level=1)
    doc.add_paragraph("在主对话框点击“规则设置…”打开折弯系数表。下图为实际运行界面。上方决定不同折弯类型使用哪一列数据，下方维护每种材料和板厚的具体数值。")
    doc.add_picture(str(rules_capture), width=Inches(6.7))
    add_rules_dialog_control_table(doc)
    add_note(doc, "规则列必须对应", "若普通折弯选择“扣除1”，程序读取当前材料和板厚行的“扣除1”；若多刀折圆选择“K因子3”，程序读取“K因子3”。不要只修改某一列，却在上方选择另一列。", fill="FFF7E6")

    doc.add_heading("6. 手动选择基面和 X 向", level=1)
    doc.add_picture(str(ASSETS / "手动选择示意.png"), width=Inches(6.7))
    add_numbered(doc, [
        "点击主对话框中的“手动选择...”。",
        "选择需要手动处理实体上的基面。",
        "程序进入 X 向边选择，选择刚才基面上的直线边。",
        "当前实体处理完成后，会继续提示选择下一个实体的基面。",
        "不再手动选择时，点击取消返回主对话框。",
        "回到主对话框点击“确定”，未手动处理的实体会继续按自动方式执行。",
    ])
    add_note(doc, "手动选择限制", "手动选择的面必须来自带有 sulian 属性且该属性有值的实体；X 向边必须是所选基面上的直线边。", fill="FFF7E6")

    doc.add_heading("7. 结果提示和失败处理", level=1)
    doc.add_picture(str(ASSETS / "结果提示.png"), width=Inches(6.7))
    add_bullets(doc, [
        "全部处理成功时，不弹最终结果框。",
        "存在失败或异常时，弹出结果明细，便于定位是哪一层、哪个部件出问题。",
        "被过滤条件跳过的对象不算失败。",
        "如果开启“展开基面颜色”，可以通过颜色检查最终使用的基面。",
        "如果选择失败改红或高亮，失败对象会按设置进行颜色或高亮标识。",
    ])

    doc.add_heading("8. 常见问题", level=1)
    qa = [
        ("为什么点“确定”后没有处理某些部件？", "先检查装配列表过滤条件、图层范围、材料/数量属性开关，以及该部件是否被勾选。"),
        ("为什么手动选择时选不到面？", "确认该实体是否有名称为 sulian 且有值的属性，并确认选择的是实体上的平面。"),
        ("为什么选不了 X 向边？", "X 向边必须在已选基面上，且必须是直线边。"),
        ("为什么全部成功没有弹框？", "这是当前设计：全部成功时不打扰用户；只有失败或异常时才弹明细。"),
        ("展开方向不对怎么办？", "对该实体使用“手动选择...”，指定正确基面和 X 向边后再执行。"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.cell(0, 0), "问题", True)
    set_cell_text(table.cell(0, 1), "处理方法", True)
    for q, a in qa:
        cells = table.add_row().cells
        set_cell_text(cells[0], q)
        set_cell_text(cells[1], a)
    style_table(table)

    doc.add_heading("9. 建议操作习惯", level=1)
    add_bullets(doc, [
        "正式批量处理前，先用少量零件验证规则和方向。",
        "复杂或方向敏感的零件优先手动指定基面和 X 向。",
        "需要复查时开启“展开基面颜色”。",
        "装配数量较多时先用过滤关键字和图层范围缩小列表。",
        "失败后先查看弹框明细，再根据图层和部件名称定位模型。",
    ])

    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("UG 智辉钣金插件 - 批量转钣金使用说明")
    r.font.name = "微软雅黑"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(120, 130, 145)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
