from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(r"D:\UGPluginRepo\ZiDonCuTu")
OUTPUT_DIR = ROOT / "outputs"
OUTPUT_DOCX = OUTPUT_DIR / "XDI-R-1000数字化曝光系统技术规格书-按附件1格式整理.docx"
OUTPUT_MD = OUTPUT_DIR / "XDI-R-1000数字化曝光系统技术规格书-按附件1格式整理.md"


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(document, text="", *, bold=False, size=12, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=Pt(6)):
    p = document.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = space_after
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    set_east_asia_font(run, "宋体")
    return p


def add_table(document, headers, rows, widths_cm=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, value in enumerate(headers):
        hdr[idx].text = str(value)
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10.5)
                set_east_asia_font(run, "宋体")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            for p in cells[idx].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    set_east_asia_font(run, "宋体")
    if widths_cm:
        for row in table.rows:
            for idx, width in enumerate(widths_cm):
                row.cells[idx].width = Cm(width)
    return table


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(10.5)
        set_east_asia_font(run, "宋体")


def build_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_paragraph(
        doc,
        "创元精密装备（拟）",
        size=12,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(80),
    )
    add_paragraph(doc, "XDI-R-1000数字化凹版曝光系统", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(10))
    add_paragraph(doc, "技术规格书", bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(10))
    add_paragraph(doc, "XDI-R-1000", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(90))
    add_paragraph(doc, "依据附件2、附件3整理，编排格式参考附件1", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(12))
    add_paragraph(doc, "2026年04月", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(0))

    doc.add_section(WD_SECTION.NEW_PAGE)
    body = doc.sections[1]
    body.top_margin = Cm(2.2)
    body.bottom_margin = Cm(2.0)
    body.left_margin = Cm(2.4)
    body.right_margin = Cm(2.2)

    add_paragraph(doc, "一、概要", bold=True, size=14)
    add_paragraph(
        doc,
        "本规格书详细说明XDI-R-1000数字化凹版曝光系统的设备用途、主要组成、关键技术指标、加工品质要求及安装环境条件。"
        "本文件根据《XDI-R-1000数字化曝光设备-2026》与《辊筒精细图形曝光设备说明》整理编制，作为方案确认、制造实施及验收对照的基础文件。",
        size=10.5,
    )

    add_paragraph(doc, "二、目的用途", bold=True, size=14)
    add_paragraph(
        doc,
        "适用于凹版制版激光-显影-蚀刻工艺中的辊筒曝光工序，通过数字化成像方式将图案直接投影至涂布有830nm红外热敏感光胶膜的辊筒表面，实现高精度图文转移。",
        size=10.5,
    )
    add_bullets(
        doc,
        [
            "适用行业：包装印刷、防伪印刷、特种材料压印。",
            "适用工件：辊筒外圆周表面，直径Φ90 mm~Φ220 mm，长度300 mm~1000 mm。",
            "加工方式：支持步进曝光与连续曝光，两端夹持支撑，辊筒可360°旋转。",
        ],
    )

    add_paragraph(doc, "三、规格概要", bold=True, size=14)
    add_paragraph(doc, "1. 设备名称型号：", bold=True, size=11)
    add_paragraph(doc, "数字化凹版曝光设备 XDI-R-1000", size=10.5)
    add_paragraph(doc, "2. 设备结构", bold=True, size=11)
    add_bullets(
        doc,
        [
            "光学引擎：830nm LED阵列、匀光系统、DMD芯片、缩小投影镜头。",
            "机械平台：花岗岩平台、内部气浮移动平台、辊筒装夹与旋转机构。",
            "曝光执行机构：10个曝光头阵列、曝光头安装横梁、X轴运动机构、六轴调节机构。",
            "调焦与控制：手动粗调+自动精调对焦机构、自动控制与图形处理软件。",
            "辅助系统：冷却系统、上下料辅助、对中装置、安全防护及环境控制系统。",
        ],
    )
    add_paragraph(doc, "3. 设备主要部件", bold=True, size=11)
    add_table(
        doc,
        ["序号", "配件", "数量", "备注"],
        [
            ["1", "光学引擎", "1套", "830nm LED阵列+匀光系统+DMD+投影镜头"],
            ["2", "曝光头阵列", "10个", "同一横梁安装，多头拼接校正"],
            ["3", "辊筒装夹与旋转机构", "1套", "适配空心辊筒/带轴辊筒"],
            ["4", "X轴运动机构", "1套", "有效行程不小于100 mm"],
            ["5", "对焦机构", "1套", "手动调焦+自动对焦"],
            ["6", "六轴调节机构", "10套", "各曝光头独立安装调节与锁紧"],
            ["7", "控制软件系统", "1套", "文件管理、曝光控制、监控报警、补偿修正"],
            ["8", "主机平台及基座", "1套", "花岗岩基座或等效高稳定性基座"],
            ["9", "冷却与温控系统", "1套", "光学系统风冷+整机温控"],
            ["10", "安全与辅助装置", "1套", "急停、门锁联锁、对中及上下料辅助"],
        ],
        widths_cm=[1.2, 4.0, 1.8, 8.0],
    )

    add_paragraph(doc, "四、设备主要部件功能介绍", bold=True, size=14)
    add_paragraph(doc, "1. 光学引擎", bold=True, size=11)
    add_table(
        doc,
        ["参数", "技术指标"],
        [
            ["工作波长", "830nm ±10nm"],
            ["投影分辨率", "12800×12800 dpi"],
            ["对应像素尺寸", "1.98μm × 1.98μm"],
            ["曝光能力", "30分钟内曝光面积至少400mm × 500mm"],
            ["光学路径", "830nm LED阵列 → 匀光系统（积分棒） → DMD芯片 → 缩小投影镜头 → 版辊表面"],
        ],
        widths_cm=[4.5, 10.5],
    )
    add_paragraph(doc, "2. 辊筒装夹与旋转机构", bold=True, size=11)
    add_paragraph(
        doc,
        "辊筒采用两端夹持支撑方式，实现360°连续旋转，并与曝光控制同步联动。设备应兼容空心辊筒与带轴辊筒两类工件。",
        size=10.5,
    )
    add_table(
        doc,
        ["辊筒类型", "直径范围", "长度范围", "特征参数"],
        [
            ["空心辊筒", "Φ90-220mm", "300-1000mm", "锥孔外径40-90mm"],
            ["带轴辊筒", "Φ90-220mm", "300-1000mm", "轴长600-1200mm"],
        ],
        widths_cm=[3.0, 3.0, 3.0, 6.0],
    )
    add_paragraph(doc, "3. 曝光头阵列与六轴调节机构", bold=True, size=11)
    add_bullets(
        doc,
        [
            "曝光头数量为10个，应安装于同一横梁或等效高刚性安装结构上。",
            "每个曝光头应具备X、Y、Z、Rx、Ry、Rz六轴调节能力，调节后可锁紧固定。",
            "Z轴需支持手动调焦与自动对焦，多曝光头之间应具备拼接校正能力。",
            "设备厂应明确各调节轴调节范围、最小分辨率、锁紧方式及锁紧后漂移量。",
        ],
    )
    add_paragraph(doc, "4. 机械运动与自动对焦", bold=True, size=11)
    add_table(
        doc,
        ["项目", "规格要求", "技术指标说明"],
        [
            ["X轴运动机构", "有效行程≥100 mm", "定位、重复定位及直线度需满足整机图形偏差要求"],
            ["辊筒旋转机构", "360°连续旋转", "旋转跳动与回转精度满足最终图形精度要求"],
            ["自动对焦", "有效分辨能力或控制精度≥0.005 mm", "适应不同辊径范围内的焦平面补偿需求"],
            ["对焦方式", "手动粗调+自动精调", "曝光过程中可调用或校准"],
        ],
        widths_cm=[3.2, 4.5, 6.8],
    )
    add_paragraph(doc, "5. 软件控制系统", bold=True, size=11)
    add_bullets(
        doc,
        [
            "支持1-bit TIFF、BMP、PNG等格式导入，并支持图形缩放、旋转、镜像、拼接处理。",
            "支持曝光控制、功率调节、时间设定、多段曝光以及步进/连续曝光模式切换。",
            "支持多曝光头同步控制、工艺参数保存调用、拼接补偿和误差修正。",
            "具备温度监测、工作状态显示、故障报警、日志记录及参数追溯功能。",
        ],
    )
    add_paragraph(doc, "6. 其他配套", bold=True, size=11)
    add_bullets(
        doc,
        [
            "主机平台建议采用花岗岩平台或等效高稳定性基座，内部采用气浮移动平台。",
            "冷却系统采用光学系统内置风冷与整机温控组合方案。",
            "配置上下料辅助装置、激光对中装置、急停按钮、红外防护罩及门锁联锁。",
        ],
    )

    add_paragraph(doc, "五、加工品质", bold=True, size=14)
    add_table(
        doc,
        ["项目", "指标", "检测方法/说明"],
        [
            ["最小图案单位（像素元）", "1.98μm × 1.98μm", "分辨率测试，确认1.98μm线宽清晰可辨"],
            ["图案变形量", "≤0.2μm/mm", "显影后用显微镜测量，1mm长度内最大偏移不超过0.2μm"],
            ["轴向直线偏移", "≤0.2μm/mm", "对1mm长轴向直线进行显微测量"],
            ["周向直线偏移", "≤0.2μm/mm", "对1mm长周向直线进行显微测量"],
            ["多头拼接与错位", "满足最终图形错位与变形要求", "设备需具备多头拼接校正能力"],
            ["径向跳动", "<0.02mm", "百分表测量辊筒外圆跳动"],
            ["重复定位精度", "±0.5μm", "激光干涉仪检测"],
            ["对焦精度", "±0.05mm", "自动对焦系统精度检测"],
            ["视场能量均匀性", ">90%", "九点法测量视场能量分布"],
            ["连续运行能力", "72小时无故障", "连续运行测试"],
        ],
        widths_cm=[4.0, 4.2, 6.3],
    )

    add_paragraph(doc, "六、放置环境条件", bold=True, size=14)
    add_table(
        doc,
        ["项目", "要求"],
        [
            ["工作环境温度", "20±1℃"],
            ["洁净度", "万级洁净度，具备防尘措施"],
            ["防震要求", "独立地基或减震垫，隔离外界振动"],
            ["上下料辅助", "气动升降台或手动叉车，便于辊筒装卸"],
            ["对中装置", "激光对心仪，确保辊筒轴线与光轴重合"],
            ["电气安全", "过载保护、漏电保护、接地保护"],
            ["激光安全", "Class 1激光安全等级，全封闭防护"],
            ["应急停机", "多处急停按钮，断电保护功能"],
            ["温度监控", "超温报警，自动停机保护"],
        ],
        widths_cm=[4.5, 10.5],
    )

    add_paragraph(doc, "备注：", bold=True, size=11)
    add_bullets(
        doc,
        [
            "本稿按附件1章节格式重组附件2、附件3内容，便于与既有规格书模板保持一致。",
            "关于设备外形尺寸、整机重量、功耗、详细安装空间等数据，附件2、附件3未给出明确数值，建议在厂家详细方案阶段补充确认。",
        ],
    )

    return doc


def build_markdown() -> str:
    return """# XDI-R-1000数字化凹版曝光系统技术规格书

副标题：按附件1格式整理

## 一、概要
本规格书详细说明XDI-R-1000数字化凹版曝光系统的设备用途、主要组成、关键技术指标、加工品质要求及安装环境条件。本文件根据附件2与附件3整理编制，作为方案确认、制造实施及验收对照的基础文件。

## 二、目的用途
适用于凹版制版激光-显影-蚀刻工艺中的辊筒曝光工序，通过数字化成像方式将图案直接投影至涂布有830nm红外热敏感光胶膜的辊筒表面，实现高精度图文转移。

- 适用行业：包装印刷、防伪印刷、特种材料压印
- 适用工件：辊筒外圆周表面，直径Φ90 mm~Φ220 mm，长度300 mm~1000 mm
- 加工方式：支持步进曝光与连续曝光，两端夹持支撑，辊筒可360°旋转

## 三、规格概要
### 1. 设备名称型号
数字化凹版曝光设备 XDI-R-1000

### 2. 设备结构
- 光学引擎：830nm LED阵列、匀光系统、DMD芯片、缩小投影镜头
- 机械平台：花岗岩平台、内部气浮移动平台、辊筒装夹与旋转机构
- 曝光执行机构：10个曝光头阵列、曝光头安装横梁、X轴运动机构、六轴调节机构
- 调焦与控制：手动粗调+自动精调对焦机构、自动控制与图形处理软件
- 辅助系统：冷却系统、上下料辅助、对中装置、安全防护及环境控制系统

### 3. 设备主要部件
| 序号 | 配件 | 数量 | 备注 |
| --- | --- | --- | --- |
| 1 | 光学引擎 | 1套 | 830nm LED阵列+匀光系统+DMD+投影镜头 |
| 2 | 曝光头阵列 | 10个 | 同一横梁安装，多头拼接校正 |
| 3 | 辊筒装夹与旋转机构 | 1套 | 适配空心辊筒/带轴辊筒 |
| 4 | X轴运动机构 | 1套 | 有效行程不小于100 mm |
| 5 | 对焦机构 | 1套 | 手动调焦+自动对焦 |
| 6 | 六轴调节机构 | 10套 | 各曝光头独立安装调节与锁紧 |
| 7 | 控制软件系统 | 1套 | 文件管理、曝光控制、监控报警、补偿修正 |
| 8 | 主机平台及基座 | 1套 | 花岗岩基座或等效高稳定性基座 |
| 9 | 冷却与温控系统 | 1套 | 光学系统风冷+整机温控 |
| 10 | 安全与辅助装置 | 1套 | 急停、门锁联锁、对中及上下料辅助 |

## 四、设备主要部件功能介绍
### 1. 光学引擎
| 参数 | 技术指标 |
| --- | --- |
| 工作波长 | 830nm ±10nm |
| 投影分辨率 | 12800×12800 dpi |
| 对应像素尺寸 | 1.98μm × 1.98μm |
| 曝光能力 | 30分钟内曝光面积至少400mm × 500mm |
| 光学路径 | 830nm LED阵列 → 匀光系统（积分棒） → DMD芯片 → 缩小投影镜头 → 版辊表面 |

### 2. 辊筒装夹与旋转机构
辊筒采用两端夹持支撑方式，实现360°连续旋转，并与曝光控制同步联动。设备应兼容空心辊筒与带轴辊筒两类工件。

| 辊筒类型 | 直径范围 | 长度范围 | 特征参数 |
| --- | --- | --- | --- |
| 空心辊筒 | Φ90-220mm | 300-1000mm | 锥孔外径40-90mm |
| 带轴辊筒 | Φ90-220mm | 300-1000mm | 轴长600-1200mm |

### 3. 曝光头阵列与六轴调节机构
- 曝光头数量为10个，应安装于同一横梁或等效高刚性安装结构上
- 每个曝光头应具备X、Y、Z、Rx、Ry、Rz六轴调节能力，调节后可锁紧固定
- Z轴需支持手动调焦与自动对焦，多曝光头之间应具备拼接校正能力
- 设备厂应明确各调节轴调节范围、最小分辨率、锁紧方式及锁紧后漂移量

### 4. 机械运动与自动对焦
| 项目 | 规格要求 | 技术指标说明 |
| --- | --- | --- |
| X轴运动机构 | 有效行程≥100 mm | 定位、重复定位及直线度需满足整机图形偏差要求 |
| 辊筒旋转机构 | 360°连续旋转 | 旋转跳动与回转精度满足最终图形精度要求 |
| 自动对焦 | 有效分辨能力或控制精度≥0.005 mm | 适应不同辊径范围内的焦平面补偿需求 |
| 对焦方式 | 手动粗调+自动精调 | 曝光过程中可调用或校准 |

### 5. 软件控制系统
- 支持1-bit TIFF、BMP、PNG等格式导入，并支持图形缩放、旋转、镜像、拼接处理
- 支持曝光控制、功率调节、时间设定、多段曝光以及步进/连续曝光模式切换
- 支持多曝光头同步控制、工艺参数保存调用、拼接补偿和误差修正
- 具备温度监测、工作状态显示、故障报警、日志记录及参数追溯功能

### 6. 其他配套
- 主机平台建议采用花岗岩平台或等效高稳定性基座，内部采用气浮移动平台
- 冷却系统采用光学系统内置风冷与整机温控组合方案
- 配置上下料辅助装置、激光对中装置、急停按钮、红外防护罩及门锁联锁

## 五、加工品质
| 项目 | 指标 | 检测方法/说明 |
| --- | --- | --- |
| 最小图案单位（像素元） | 1.98μm × 1.98μm | 分辨率测试，确认1.98μm线宽清晰可辨 |
| 图案变形量 | ≤0.2μm/mm | 显影后用显微镜测量，1mm长度内最大偏移不超过0.2μm |
| 轴向直线偏移 | ≤0.2μm/mm | 对1mm长轴向直线进行显微测量 |
| 周向直线偏移 | ≤0.2μm/mm | 对1mm长周向直线进行显微测量 |
| 多头拼接与错位 | 满足最终图形错位与变形要求 | 设备需具备多头拼接校正能力 |
| 径向跳动 | <0.02mm | 百分表测量辊筒外圆跳动 |
| 重复定位精度 | ±0.5μm | 激光干涉仪检测 |
| 对焦精度 | ±0.05mm | 自动对焦系统精度检测 |
| 视场能量均匀性 | >90% | 九点法测量视场能量分布 |
| 连续运行能力 | 72小时无故障 | 连续运行测试 |

## 六、放置环境条件
| 项目 | 要求 |
| --- | --- |
| 工作环境温度 | 20±1℃ |
| 洁净度 | 万级洁净度，具备防尘措施 |
| 防震要求 | 独立地基或减震垫，隔离外界振动 |
| 上下料辅助 | 气动升降台或手动叉车，便于辊筒装卸 |
| 对中装置 | 激光对心仪，确保辊筒轴线与光轴重合 |
| 电气安全 | 过载保护、漏电保护、接地保护 |
| 激光安全 | Class 1激光安全等级，全封闭防护 |
| 应急停机 | 多处急停按钮，断电保护功能 |
| 温度监控 | 超温报警，自动停机保护 |

## 备注
- 本稿按附件1章节格式重组附件2、附件3内容，便于与既有规格书模板保持一致
- 关于设备外形尺寸、整机重量、功耗、详细安装空间等数据，附件2、附件3未给出明确数值，建议在厂家详细方案阶段补充确认
"""


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT_DOCX)
    OUTPUT_MD.write_text(build_markdown(), encoding="utf-8")
    print(OUTPUT_DOCX)
    print(OUTPUT_MD)


if __name__ == "__main__":
    main()
